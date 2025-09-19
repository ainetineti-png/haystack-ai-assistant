from fastapi import FastAPI, Request
from fastapi.middleware.cors import CORSMiddleware
import os
from typing import List, Dict, Any, Set, Optional  # Added Optional
import json
import requests
import glob
import PyPDF2
import pdfplumber
from docx import Document
import uuid  # Added for valid UUID point ids
import hashlib  # For chunk content hashing
from contextlib import asynccontextmanager  # re-add lifespan decorator after earlier edit
import atexit  # Safe shutdown handler
import unicodedata  # For block normalization
import statistics  # For table summarization stats
from collections import OrderedDict  # Added for rerank LRU cache
from collections import deque  # NEW: for conversation history
import re  # Added for prompt / citation post-processing
# Removed LangExtract import - no longer needed
from qdrant_client import QdrantClient, models as qdrant_models
from sentence_transformers import SentenceTransformer
import sys
from pathlib import Path
import time  # For timing operations
from sqlalchemy import create_engine, Column, Integer, String, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from datetime import datetime
from pydantic import BaseModel

# Load environment variables from config.env if it exists
def load_env_from_file(env_file):
    if os.path.exists(env_file):
        print(f"Loading environment from {env_file}")
        with open(env_file, 'r') as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith('#'):
                    continue
                key, value = line.split('=', 1)
                os.environ[key] = value

# Try to load from different possible locations
config_paths = [
    os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'config.env'),  # Project root
    os.path.join(os.path.dirname(os.path.abspath(__file__)), 'config.env'),  # Backend folder
    os.path.join(os.getcwd(), 'config.env'),  # Current working directory
    os.path.join(Path(sys.executable).parent, 'config.env')  # Executable directory (for PyInstaller)
]

for config_path in config_paths:
    if os.path.exists(config_path):
        load_env_from_file(config_path)
        break
from threading import Lock
import numpy as np  # ensure np for similarity array handling
import pickle
from scipy import sparse as _scipy_sparse  # for persistence of sparse matrix

# ---------------- New: Reload / Supervisor Detection ---------------- #
IS_RELOAD_SUPERVISOR = os.environ.get('WATCHGOD_RELOADER') == 'true'
# If running under uvicorn --reload and using embedded Qdrant, concurrent access will fail.
# We will skip heavy init in supervisor and optionally auto-disable reload for worker if embedded path in use.
FORCE_DISABLE_RELOAD_FOR_EMBEDDED = os.environ.get('FORCE_DISABLE_RELOAD_FOR_EMBEDDED', '1') == '1'
# --------------------------------------------------------------- #

# ---------------- New: Deferred heavy init flag ---------------- #
DEFER_STARTUP_INIT = os.environ.get('DEFER_STARTUP_INIT', '1') == '1'
_STARTUP_INITIALIZED = False
# --------------------------------------------------------------- #

# ---------------- New: BM25 Globals ---------------- #
SPARSE_METHOD = os.environ.get('SPARSE_METHOD', 'tfidf').lower()  # 'tfidf' or 'bm25'
_bm25_corpus_tokens: List[List[str]] = []
_bm25_doc_freq: Dict[str, int] = {}
_bm25_inverted_index: Dict[str, List[tuple]] = {}
_bm25_doc_len: List[int] = []
_bm25_avgdl: float = 0.0
_bm25_k1 = float(os.environ.get('BM25_K1', '1.5'))
_bm25_b = float(os.environ.get('BM25_B', '0.75'))
# --------------------------------------------------- #

# Sparse / rerank imports
try:
    from sklearn.feature_extraction.text import TfidfVectorizer  # type: ignore
    from sklearn.metrics.pairwise import cosine_similarity  # type: ignore
    _SPARSE_OK = True
except Exception:
    TfidfVectorizer = None  # type: ignore
    cosine_similarity = None  # type: ignore
    _SPARSE_OK = False
try:
    from sentence_transformers import CrossEncoder
    _CROSS_ENCODER_MODEL_NAME = os.environ.get('CROSS_ENCODER_MODEL', 'cross-encoder/ms-marco-MiniLM-L-6-v2')
    _ENABLE_RERANK = os.environ.get('ENABLE_RERANK', '0') == '1'
    cross_encoder = CrossEncoder(_CROSS_ENCODER_MODEL_NAME) if _ENABLE_RERANK else None
except Exception:
    cross_encoder = None
    _ENABLE_RERANK = False

ENABLE_SPARSE = os.environ.get('ENABLE_SPARSE', '1') == '1'
SPARSE_MAX_DOCS = int(os.environ.get('SPARSE_MAX_DOCS', '50000'))
RRF_K = int(os.environ.get('RRF_K', '60'))
TABLE_SUMMARY_MAX_ROWS = int(os.environ.get('TABLE_SUMMARY_MAX_ROWS', '8'))
TABLE_SUMMARY_MAX_COLS = int(os.environ.get('TABLE_SUMMARY_MAX_COLS', '8'))
TABLE_CSV_MAX_CHARS = int(os.environ.get('TABLE_CSV_MAX_CHARS', '20000'))  # New: cap stored CSV size
ENABLE_SPARSE_PERSIST = os.environ.get('ENABLE_SPARSE_PERSIST', '1') == '1'
SPARSE_PERSIST_DIR = os.environ.get('SPARSE_PERSIST_DIR', os.path.join(os.path.dirname(__file__), 'sparse_index'))
# Rerank caching controls
RERANK_CACHE_SIZE = int(os.environ.get('RERANK_CACHE_SIZE', '5000'))
RERANK_TOP_N = int(os.environ.get('RERANK_TOP_N', '50'))  # max candidates to send to cross-encoder
# --------------------------------------------------- #

# New: context assembly / diversity flags
MAX_CONTEXT_TOKENS = int(os.environ.get('MAX_CONTEXT_TOKENS', '1800'))
ENABLE_DIVERSITY = os.environ.get('ENABLE_DIVERSITY', '1') == '1'
TABLE_KEYWORDS = [
    'table', 'tabulate', 'dataset', 'results', 'values', 'fig.', 'figure', 'data', 'statistics'
]

_sparse_lock = Lock()
_sparse_vectorizer = None
_sparse_matrix = None
_sparse_payload_refs: List[dict] = []  # store minimal chunk refs
# Rerank LRU cache: key=(query_md5 + chunk_md5) -> score
_rerank_cache: 'OrderedDict[str,float]' = OrderedDict()

# Cross-encoder rerank helper (LRU cache)

def _apply_cross_encoder_rerank(query: str, fused: List[dict]):
    if not fused or cross_encoder is None:
        return
    candidates = fused[:RERANK_TOP_N]
    q_md5 = hashlib.md5(query.encode('utf-8')).hexdigest()
    pairs = []
    idxs = []
    for i, item in enumerate(candidates):
        chunk_hash = item.get('md5') or item.get('id') or str(i)
        cache_key = f"{q_md5}|{chunk_hash}"
        item['_rerank_cache_key'] = cache_key
        cached = _rerank_cache.get(cache_key)
        if cached is not None:
            item['rerank_score'] = cached
            _rerank_cache.move_to_end(cache_key)
            continue
        text = item.get('content') or item.get('text') or ''
        if not text:
            item['rerank_score'] = 0.0
            continue
        pairs.append((query, text))
        idxs.append(i)
    if pairs:
        try:
            scores = cross_encoder.predict(pairs)
            for local_i, score in zip(idxs, scores):
                cand = candidates[local_i]
                cand['rerank_score'] = float(score)
                ck = cand.get('_rerank_cache_key')
                if ck:
                    _rerank_cache[ck] = float(score)
        except Exception as e:
            print(f"[WARNING] Cross-encoder prediction failed: {e}")
    while len(_rerank_cache) > RERANK_CACHE_SIZE:
        try:
            _rerank_cache.popitem(last=False)
        except Exception:
            break
    for item in candidates:
        if 'rerank_score' not in item:
            item['rerank_score'] = 0.0
    fused.sort(key=lambda x: (-(x.get('rerank_score', 0.0)), -(x.get('rrf_score', 0.0))))

# ---------------- Sparse Persistence Helpers ---------------- #

def _persist_sparse_index():
    if not ENABLE_SPARSE_PERSIST:
        return
    # TF-IDF persistence
    try:
        if _sparse_vectorizer is not None and _sparse_matrix is not None:
            os.makedirs(SPARSE_PERSIST_DIR, exist_ok=True)
            with open(os.path.join(SPARSE_PERSIST_DIR, 'vectorizer.pkl'), 'wb') as vf:
                pickle.dump(_sparse_vectorizer, vf)
            sparse_path = os.path.join(SPARSE_PERSIST_DIR, 'matrix.npz')
            _scipy_sparse.save_npz(sparse_path, _sparse_matrix)
            with open(os.path.join(SPARSE_PERSIST_DIR, 'payload_refs.json'), 'w', encoding='utf-8') as pf:
                json.dump(_sparse_payload_refs, pf)
    except Exception as e:
        print(f"[WARNING] Failed to persist TF-IDF sparse index: {e}")
    # BM25 persistence
    try:
        if SPARSE_METHOD == 'bm25' and _bm25_corpus_tokens:
            os.makedirs(SPARSE_PERSIST_DIR, exist_ok=True)
            bm25_state = {
                'corpus_tokens': _bm25_corpus_tokens,
                'doc_freq': _bm25_doc_freq,
                'inverted_index': _bm25_inverted_index,
                'doc_len': _bm25_doc_len,
                'avgdl': _bm25_avgdl,
                'payload_refs': _sparse_payload_refs,
                'k1': _bm25_k1,
                'b': _bm25_b
            }
            with open(os.path.join(SPARSE_PERSIST_DIR, 'bm25_index.pkl'), 'wb') as bf:
                pickle.dump(bm25_state, bf)
    except Exception as e:
        print(f"[WARNING] Failed to persist BM25 index: {e}")
    try:
        print(f"[INFO] Persisted sparse index to {SPARSE_PERSIST_DIR}")
    except Exception:
        pass

def _load_sparse_index_from_disk():
    global _sparse_vectorizer, _sparse_matrix, _sparse_payload_refs
    global _bm25_corpus_tokens, _bm25_doc_freq, _bm25_inverted_index, _bm25_doc_len, _bm25_avgdl
    if not ENABLE_SPARSE_PERSIST:
        return
    # Load TF-IDF if present (even if SPARSE_METHOD is bm25, for debugging)
    try:
        vec_path = os.path.join(SPARSE_PERSIST_DIR, 'vectorizer.pkl')
        mat_path = os.path.join(SPARSE_PERSIST_DIR, 'matrix.npz')
        refs_path = os.path.join(SPARSE_PERSIST_DIR, 'payload_refs.json')
        if os.path.isfile(vec_path) and os.path.isfile(mat_path) and os.path.isfile(refs_path):
            with open(vec_path, 'rb') as vf:
                _sparse_vectorizer = pickle.load(vf)
            _sparse_matrix = _scipy_sparse.load_npz(mat_path)
            with open(refs_path, 'r', encoding='utf-8') as pf:
                _sparse_payload_refs = json.load(pf)
            print(f"[INFO] Loaded persisted TF-IDF index: docs={len(_sparse_payload_refs)} features={_sparse_matrix.shape[1] if _sparse_matrix is not None else 0}")
    except Exception as e:
        print(f"[WARNING] Could not load persisted TF-IDF index: {e}")
    # Load BM25 if requested
    if SPARSE_METHOD == 'bm25':
        try:
            bm25_path = os.path.join(SPARSE_PERSIST_DIR, 'bm25_index.pkl')
            if os.path.isfile(bm25_path):
                with open(bm25_path, 'rb') as bf:
                    state = pickle.load(bf)
                _bm25_corpus_tokens = state.get('corpus_tokens', [])
                _bm25_doc_freq = state.get('doc_freq', {})
                _bm25_inverted_index = state.get('inverted_index', {})
                _bm25_doc_len = state.get('doc_len', [])
                _bm25_avgdl = state.get('avgdl', 0.0)
                _sparse_payload_refs = state.get('payload_refs', [])  # reuse refs
                print(f"[INFO] Loaded persisted BM25 index: docs={len(_bm25_corpus_tokens)} avgdl={_bm25_avgdl:.2f}")
        except Exception as e:
            print(f"[WARNING] Could not load persisted BM25 index: {e}")
# ---------------- End Sparse Persistence Helpers ---------------- #

# Define lifespan before app instantiation (ensure AVAILABLE global symbols referenced inside) 
@asynccontextmanager
async def lifespan(app: FastAPI):
    global _STARTUP_INITIALIZED, embedder, qdrant_client
    # Perform heavy initialization only once at real server startup (not in reloader supervisor)
    if DEFER_STARTUP_INIT and not _STARTUP_INITIALIZED:
        try:
            # Initialize Qdrant client
            if qdrant_client is None:
                _initialize_qdrant_client()
            # Load embedding model lazily
            if embedder is None:
                default_model_path_local = os.path.join(os.path.dirname(os.path.dirname(__file__)), "models", "all-MiniLM-L6-v2")
                model_path_local = os.environ.get("EMBED_MODEL_PATH", default_model_path_local)
                print(f"[INIT] Deferred model load from: {model_path_local}")
                try:
                    embedder = SentenceTransformer(model_path_local)
                    print(f"[INIT] Successfully loaded embedding model (deferred) from: {model_path_local}")
                except Exception as e:
                    print(f"[WARNING] Deferred model load failed: {e}")
                    embedder = None
            _STARTUP_INITIALIZED = True
        except Exception as e:
            print(f"[WARNING] Deferred heavy init encountered error: {e}")
    # Only ingest after heavy init to ensure embedder/qdrant available
    if DEFER_STARTUP_INIT:
        if 'AUTO_INGEST' in globals() and globals().get('AUTO_INGEST'):
            try:
                load_documents()
            except Exception as e:
                print(f"[WARNING] Initial load_documents failed (deferred): {e}")
        else:
            print("[INFO] AUTO_INGEST disabled; skipping ingestion at startup (deferred)")
    else:
        # Original path (no defer) retains previous behavior
        if 'AUTO_INGEST' in globals() and globals().get('AUTO_INGEST'):
            try:
                load_documents()
            except Exception as e:
                print(f"[WARNING] Initial load_documents failed: {e}")
        else:
            print("[INFO] AUTO_INGEST disabled; skipping ingestion at startup")
    # Attempt to load persisted sparse index after initial load
    _load_sparse_index_from_disk()
    try:
        yield
    finally:
        try:
            _safe_close_qdrant()
        except Exception:
            pass

# ---------------- Sparse / Hybrid Retrieval Helpers (Phase 3) ---------------- #

def _tokenize(text: str) -> List[str]:
    return [t for t in text.lower().split() if t]

def _build_sparse_index():
    # If vectorizer class unavailable and method tfidf, skip
    if SPARSE_METHOD == 'tfidf' and TfidfVectorizer is None:
        return
    """Build a sparse index over all chunk texts currently in Qdrant.
    Supports TF-IDF (default) or BM25 based on SPARSE_METHOD env variable.
    Thread-safe under _sparse_lock."""
    global _sparse_vectorizer, _sparse_matrix, _sparse_payload_refs
    global _bm25_corpus_tokens, _bm25_doc_freq, _bm25_inverted_index, _bm25_doc_len, _bm25_avgdl
    if not (ENABLE_SPARSE and _SPARSE_OK and embedder is not None):
        # Allow BM25 even if scikit is missing (we don't depend on sklearn for BM25)
        if SPARSE_METHOD != 'bm25':
            return
    if qdrant_client is None:
        print("[INFO] Qdrant client not available - skipping sparse index build")
        return
    try:
        with _sparse_lock:
            texts: List[str] = []
            refs: List[dict] = []
            if QDRANT_COLLECTION not in [c.name for c in qdrant_client.get_collections().collections]:
                print("[INFO] Qdrant collection not found - skipping sparse index build")
                return
            offset = None
            batch = 512
            count = 0
            while True:
                points, offset = qdrant_client.scroll(collection_name=QDRANT_COLLECTION, scroll_filter=None, limit=batch, with_payload=True, offset=offset)
                if not points:
                    break
                for pt in points:
                    if count >= SPARSE_MAX_DOCS:
                        break
                    payload = pt.payload or {}
                    text = payload.get('text')
                    if not text:
                        continue
                    texts.append(text)
                    refs.append({
                        'id': str(pt.id),
                        'filename': payload.get('filename'),
                        'page': payload.get('page'),
                        'chunk_index': payload.get('chunk_index'),
                        'heading_path': payload.get('heading_path'),
                        'block_types': payload.get('block_types'),
                        'md5': payload.get('md5'),  # added for rerank caching key
                        'text': text
                    })
                    count += 1
                if count >= SPARSE_MAX_DOCS or offset is None:
                    break
            if not texts:
                print("[INFO] No texts found for sparse index build")
                return
            _sparse_payload_refs = refs
            if SPARSE_METHOD == 'tfidf':
                if TfidfVectorizer is None:  # Defensive guard for type checker
                    return
                try:
                    _sparse_vectorizer = TfidfVectorizer(max_df=0.9, min_df=1, ngram_range=(1,2))
                    _sparse_matrix = _sparse_vectorizer.fit_transform(texts)
                    print(f"[INFO] Built TF-IDF index: docs={len(refs)} features={_sparse_matrix.shape[1]}")
                except Exception as ve:
                    print(f"[WARNING] TF-IDF vectorizer build failed: {ve}")
                    return
            else:  # BM25
                _bm25_corpus_tokens = []
                _bm25_doc_freq = {}
                _bm25_inverted_index = {}
                _bm25_doc_len = []
                for doc_idx, txt in enumerate(texts):
                    tokens = _tokenize(txt)
                    _bm25_corpus_tokens.append(tokens)
                    _bm25_doc_len.append(len(tokens))
                    tf_local: Dict[str, int] = {}
                    for t in tokens:
                        tf_local[t] = tf_local.get(t, 0) + 1
                    for term, tf in tf_local.items():
                        _bm25_doc_freq[term] = _bm25_doc_freq.get(term, 0) + 1
                        _bm25_inverted_index.setdefault(term, []).append((doc_idx, tf))
                N = len(_bm25_corpus_tokens)
                _bm25_avgdl = sum(_bm25_doc_len)/N if N else 0.0
                print(f"[INFO] Built BM25 index: docs={N} avgdl={_bm25_avgdl:.2f} vocab={len(_bm25_doc_freq)}")
            _persist_sparse_index()
    except Exception as e:
        print(f"[WARNING] Failed to build sparse index: {e}")


def _sparse_search(query: str, top_k: int) -> List[dict]:
    if SPARSE_METHOD == 'bm25':
        return _bm25_search(query, top_k)
    # TF-IDF path (existing)
    if TfidfVectorizer is None or cosine_similarity is None:
        return []
    try:
        q_vec = _sparse_vectorizer.transform([query]) if _sparse_vectorizer is not None else None
        if q_vec is None or _sparse_matrix is None:
            return []
        sims_arr = cosine_similarity(q_vec, _sparse_matrix).ravel() if cosine_similarity is not None else None
        if sims_arr is None or getattr(sims_arr, 'size', 0) == 0:
            return []
        sims = np.asarray(sims_arr)
        top_idx = sims.argsort()[::-1][:top_k]
        results = []
        for rank, idx in enumerate(top_idx):
            ref = _sparse_payload_refs[idx]
            results.append({
                'id': ref['id'],
                'filename': ref['filename'],
                'page': ref['page'],
                'chunk_index': ref['chunk_index'],
                'heading_path': ref.get('heading_path'),
                'block_types': ref.get('block_types'),
                'content': ref['text'],
                'md5': ref.get('md5'),  # propagate md5
                'score': float(sims[idx]),
                'source': 'sparse_tfidf',
                'rank': rank+1
            })
        return results
    except Exception as e:
        print(f"[WARNING] Sparse search failed: {e}")
        return []


def _bm25_search(query: str, top_k: int) -> List[dict]:
    if SPARSE_METHOD != 'bm25' or not _bm25_corpus_tokens:
        return []
    try:
        q_terms = _tokenize(query)
        if not q_terms:
            return []
        scores: Dict[int, float] = {}
        N = len(_bm25_corpus_tokens)
        avgdl = _bm25_avgdl if _bm25_avgdl else 1.0
        k1 = _bm25_k1
        b = _bm25_b
        seen_docs: Set[int] = set()
        for qt in q_terms:
            postings = _bm25_inverted_index.get(qt)
            if not postings:
                continue
            df = _bm25_doc_freq.get(qt, 0)
            # BM25 IDF (Robertson/Sparck Jones)
            idf = np.log(1 + (N - df + 0.5)/(df + 0.5)) if df else 0.0
            for doc_idx, tf in postings:
                dl = _bm25_doc_len[doc_idx] or 1
                denom = tf + k1 * (1 - b + b * dl / avgdl)
                score_add = idf * (tf * (k1 + 1)) / denom if denom else 0.0
                scores[doc_idx] = scores.get(doc_idx, 0.0) + score_add
                seen_docs.add(doc_idx)
        if not scores:
            return []
        ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)[:top_k]
        results: List[dict] = []
        for rank, (doc_idx, sc) in enumerate(ranked, start=1):
            ref = _sparse_payload_refs[doc_idx]
            results.append({
                'id': ref['id'],
                'filename': ref['filename'],
                'page': ref['page'],
                'chunk_index': ref['chunk_index'],
                'heading_path': ref.get('heading_path'),
                'block_types': ref.get('block_types'),
                'content': ref['text'],
                'md5': ref.get('md5'),  # propagate md5
                'score': float(sc),
                'source': 'sparse_bm25',
                'rank': rank
            })
        return results
    except Exception as e:
        print(f"[WARNING] BM25 search failed: {e}")
        return []


def _dense_search_internal(query: str, top_k: int) -> List[dict]:
    if embedder is None or qdrant_client is None:
        return []
    try:
        query_vec = embedder.encode(query).tolist()
        try:
            qp = qdrant_client.query_points(collection_name=QDRANT_COLLECTION, query=query_vec, limit=top_k, with_payload=True)
            hits = qp.points
        except Exception:
            hits = qdrant_client.search(collection_name=QDRANT_COLLECTION, query_vector=query_vec, limit=top_k, with_payload=True)
        out = []
        for rank, hit in enumerate(hits):
            payload = hit.payload or {}
            out.append({
                'id': str(hit.id),
                'filename': payload.get('filename'),
                'page': payload.get('page'),
                'chunk_index': payload.get('chunk_index'),
                'heading_path': payload.get('heading'),
                'block_types': payload.get('block_types'),
                'content': payload.get('text', ''),
                'md5': payload.get('md5'),  # for rerank cache
                'score': float(hit.score) if hasattr(hit, 'score') else 0.0,
                'source': 'dense',
                'rank': rank+1
            })
        return out
    except Exception as e:
        print(f"[WARNING] Dense search failed: {e}")
        return []


def _rrf_fuse(dense: List[dict], sparse: List[dict], final_k: int) -> List[dict]:
    """Reciprocal Rank Fusion over two ranked lists."""
    fused: Dict[str, dict] = {}
    for lst in (dense, sparse):
        for item in lst:
            key = item['id']
            rank = item['rank']
            contrib = 1.0 / (RRF_K + rank)
            if key not in fused:
                fused[key] = {**item, 'rrf_score': contrib}
            else:
                fused[key]['rrf_score'] += contrib
    fused_list = list(fused.values())
    fused_list.sort(key=lambda x: (-x['rrf_score'], x['filename'] or '', x['page'] or 0))
    return fused_list[:final_k]


def hybrid_search(query: str, top_k_dense: int = 20, top_k_sparse: int = 20, final_k: int = 5) -> List[dict]:
    dense_hits = _dense_search_internal(query, top_k_dense) if embedder is not None else []
    sparse_hits = _sparse_search(query, top_k_sparse)
    if not dense_hits and not sparse_hits:
        return []
    fused = _rrf_fuse(dense_hits, sparse_hits, final_k * 3)
    # Optional cross-encoder rerank with caching
    if cross_encoder is not None and _ENABLE_RERANK and fused:
        try:
            _apply_cross_encoder_rerank(query, fused)
        except Exception as e:
            print(f"[WARNING] Cross-encoder rerank failed: {e}")
    final = fused[:final_k]
    return [
        {
            'filename': c['filename'],
            'content': c['content'],
            'page': c['page'],
            'chunk_index': c['chunk_index'],
            'heading_path': c.get('heading_path'),
            'block_types': c.get('block_types'),
            'md5': c.get('md5')
        } for c in final
    ]
# ---------------- End Sparse / Hybrid Retrieval Helpers ---------------- #

# ---------------- Context Assembly & Generation Engine (Added) ---------------- #
# Lightweight in-file implementations to fix NameError for assemble_context & generate_answer_engine
# If later modularizing, these can be moved to a dedicated module.

from typing import Tuple  # ensure Tuple available
import re, time
try:
    import requests  # for Ollama calls
except Exception:
    requests = None  # type: ignore

# In-memory conversation history (session_id -> list of {question, answer})
_conversation_histories: Dict[str, List[dict]] = {}
_MAX_HISTORY = 8  # keep last N turns

def assemble_context(question: str, raw_docs: List[dict], max_tokens: int = 1800) -> List[dict]:
    """Assemble a trimmed context list respecting a rough token budget.
    Currently uses a simple word-count approximation (1 word ~ 1 token for short English).
    Preserves original ordering from retrieval. Assigns provisional citation ids later in generation.
    """
    approx_tokens = 0
    context: List[dict] = []
    for d in raw_docs:
        text = d.get('content') or d.get('text') or ''
        # Trim excessively long chunks to avoid blowing the budget
        if len(text) > 1200:
            text = text[:1200]
        token_est = len(text.split())
        if context and approx_tokens + token_est > max_tokens:
            break
        nd = dict(d)  # shallow copy
        nd['content'] = text
        context.append(nd)
        approx_tokens += token_est
    return context

def prompt_builder(question: str, context_docs: List[dict], chat_summary: Optional[str] = None, mode: str = "tutor", history: Optional[List[dict]] = None) -> str:
    system_prompt = {
        "tutor": """You are a patient, expert tutor. ALWAYS structure your answer in this exact format using Markdown:

# 🎯 Problem
[One clear sentence restating the question]

## 📚 Key Concepts
- [Bullet point 1: Core concept/formula needed]
- [Bullet point 2: Another key idea]  
- [Bullet point 3: etc.]

## 📝 Step-by-Step Solution
1. **Step 1**: [Clear action] - [Brief explanation]
2. **Step 2**: [Next action] - [Show any equations/calculations]
3. **Step 3**: [Continue logically] - [Explain reasoning]
4. **Final Step**: [Conclude] - [State the answer clearly]

## ✅ Final Answer
**[State the final answer clearly in bold or use `code formatting` for numbers]**

## 🔍 Quick Check
[One sentence verification or dimensional analysis to confirm the answer makes sense]

## 📖 References
[List sources you actually used, format: [C1] filename page X]

Rules:
- Use citations [C1], [C2] ONLY where you actually reference the provided context
- Keep each step concise but complete
- Use proper Markdown formatting with headers, bold, bullets
- If using general knowledge not in context, say "(general physics knowledge)" 
- Be encouraging and educational, not just computational""",
        "citations": "You are an academic assistant. Use ONLY the provided context. Cite facts with [C1], [C2], etc. Include a References section in Markdown format.",
        "concise": "You are a helpful assistant. Answer concisely using the provided context in well-formatted Markdown."
    }.get(mode, "You are a helpful assistant.")
    parts = [system_prompt]
    if chat_summary:
        parts.append(f"Chat summary: {chat_summary}")
    if history:
        for h in history[-3:]:
            parts.append(f"Prev Q: {h.get('question','')[:140]}\nA: {h.get('answer','')[:200]}")
    ctx_lines = []
    for i, d in enumerate(context_docs, start=1):
        snippet = (d.get('content','') or '')[:300].strip()
        ctx_lines.append(f"[C{i}] Source: {d.get('filename')} page {d.get('page')} chunk {d.get('chunk_index')}\n{snippet}")
    parts.append("Question: " + question)
    parts.append("Context:\n" + "\n\n".join(ctx_lines))
    parts.append("Answer:")
    return "\n\n".join(parts)

def postprocess_answer(answer: str, context_docs: List[dict], question: str = "", min_words: int = 20, max_words: int = 500) -> str:
    # Normalize citation markers to [C#]
    answer = re.sub(r"\[C\s*(\d+)\]", r"[C\1]", answer)
    answer = re.sub(r"\(C(\d+)\)", r"[C\1]", answer)
    valid = {str(i+1) for i in range(len(context_docs))}
    answer = re.sub(r"\[C(\d+)\]", lambda m: m.group(0) if m.group(1) in valid else "", answer)
    
    # Ensure required sections exist
    required_sections = ["Problem", "Key Concepts", "Step-by-Step Solution", "Final Answer", "Quick Check", "References"]
    for section in required_sections:
        if f"# {section}" not in answer and f"## {section}" not in answer:
            if section == "Problem" and "🎯 Problem" not in answer:
                problem_text = question if question else "Problem analysis needed"
                answer = f"# 🎯 Problem\n{problem_text}\n\n" + answer
            elif section == "Key Concepts" and "📚 Key Concepts" not in answer:
                answer += f"\n\n## 📚 Key Concepts\n- Core concepts from the provided context"
            elif section == "Step-by-Step Solution" and "📝 Step-by-Step Solution" not in answer:
                answer += f"\n\n## 📝 Step-by-Step Solution\n1. **Analysis**: Based on the given information"
            elif section == "Final Answer" and "✅ Final Answer" not in answer:
                answer += f"\n\n## ✅ Final Answer\n**Answer derived from the analysis above**"
            elif section == "Quick Check" and "🔍 Quick Check" not in answer:
                answer += f"\n\n## 🔍 Quick Check\nThe answer appears reasonable based on the given parameters."
    
    # Ensure References section with actual citations
    if "References" not in answer and "📖 References" not in answer and context_docs:
        refs = "\n\n## 📖 References\n" + "\n".join([f"[C{i+1}] {d.get('filename')} page {d.get('page')}" for i, d in enumerate(context_docs[:3])])
        answer += refs
    
    # Length control (increased limit for better educational content)
    words = answer.split()
    if len(words) < min_words:
        answer += "\n\n*Note: Answer enhanced for educational completeness.*"
    elif len(words) > max_words:
        # Try to trim while keeping structure
        lines = answer.split('\n')
        truncated_lines = []
        word_count = 0
        for line in lines:
            line_words = len(line.split())
            if word_count + line_words <= max_words:
                truncated_lines.append(line)
                word_count += line_words
            elif line.startswith('#') or line.startswith('##'):  # Keep headers
                truncated_lines.append(line)
            else:
                break
        answer = '\n'.join(truncated_lines)
        if word_count >= max_words:
            answer += "\n\n*[Answer truncated for length]*"
    
    return answer

def _detect_topic_shift(history: List[dict], new_question: str) -> bool:
    if not history:
        return True
    last_q = history[-1].get('question','')
    if not last_q:
        return True
    q_tokens = set(new_question.lower().split())
    last_tokens = set(last_q.lower().split())
    if not last_tokens:
        return True
    overlap = len(q_tokens & last_tokens) / max(1, len(q_tokens | last_tokens))
    return overlap < 0.25  # heuristic threshold

def generate_answer_engine(question: str, context_docs: List[dict], chat_summary: Optional[str] = None, session_id: Optional[str] = None, mode: str = "tutor") -> Tuple[str, List[dict], Dict[str, Any]]:
    """Unified answer generation + metadata.
    Returns (answer, citations, meta) where meta includes topic_shift & history_used.
    """
    # Assign citation ids
    for idx, d in enumerate(context_docs, start=1):
        d['citation_id'] = idx
    history = _conversation_histories.get(session_id, []) if session_id else []
    topic_shift = _detect_topic_shift(history, question)
    history_used = bool(history)
    prompt = prompt_builder(question, context_docs, chat_summary, mode, history if history_used else None)

    # Call model (Ollama) with retries
    answer = ""
    retries = 3
    initial_timeout = 80
    min_timeout = 25
    timeout_reduction = 0.7
    timeout_count = 0
    if requests is None:
        answer = "Error: requests library unavailable to contact model."
    else:
        for attempt in range(retries + 1):
            current_timeout = max(min_timeout, initial_timeout * (timeout_reduction ** attempt))
            try:
                resp = requests.post(
                    "http://localhost:11434/api/generate",
                    json={
                        "model": os.environ.get('OLLAMA_MODEL','llama3:latest'),
                        "prompt": prompt,
                        "stream": False,
                        "temperature": 0.7
                    },
                    timeout=current_timeout
                )
                resp.raise_for_status()
                data = resp.json()
                answer = data.get('response','').strip()
                if not answer and attempt < retries:
                    timeout_count += 1
                    time.sleep(2 * (attempt+1))
                    continue
                if not answer:
                    answer = "Error: Empty response from model."
                break
            except requests.exceptions.Timeout:
                timeout_count += 1
                if attempt == retries:
                    answer = f"Error: Model timeout after {retries+1} attempts."
                else:
                    time.sleep(2 * (attempt+1))
            except requests.exceptions.ConnectionError:
                if attempt == retries:
                    answer = "Error: Cannot connect to model server (Ollama)."
                else:
                    time.sleep(3 * (attempt+1))
            except Exception as e:
                if attempt == retries:
                    answer = f"Error: Model request failed: {e}"[:400]
                else:
                    time.sleep(2 * (attempt+1))
                    timeout_count += 1
    answer = postprocess_answer(answer, context_docs, question)
    citations = [
        {
            'id': d['citation_id'],
            'filename': d.get('filename'),
            'page': d.get('page'),
            'chunk_index': d.get('chunk_index')
        }
        for d in context_docs[: max(3, min(len(context_docs), 6) - timeout_count)]
    ]
    # Update history
    if session_id:
        hist = _conversation_histories.setdefault(session_id, [])
        hist.append({'question': question, 'answer': answer[:800]})
        if len(hist) > _MAX_HISTORY:
            del hist[0: len(hist) - _MAX_HISTORY]
    meta = {
        'topic_shift': topic_shift,
        'history_used': history_used,
        'retries': retries,
        'timeout_count': timeout_count,
        'context_docs': len(context_docs)
    }
    return answer, citations, meta
# ---------------- End Context Assembly & Generation Engine (Added) ---------------- #

# Instantiate app here after helper definitions (lifespan previously defined above)
app = FastAPI(lifespan=lifespan)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Use environment variables with fallbacks for paths
DATA_DIR = os.environ.get("DATA_DIR", os.path.join(os.path.dirname(__file__), "data"))
documents = []
last_loaded_files = {}

# Progress tracking globals
ingest_total_files = 0
ingest_processed_files = 0

# Qdrant setup (embedded/local mode)
# Use environment variable with fallback for Qdrant path
QDRANT_PATH = os.environ.get("QDRANT_PATH", os.path.join(os.path.dirname(__file__), "qdrant_storage"))
QDRANT_COLLECTION = "documents"
print(f"Using Qdrant storage path: {QDRANT_PATH}" if not IS_RELOAD_SUPERVISOR else "[SUPERVISOR] Detected reloader supervisor; deferring embedded Qdrant init")

# Create Qdrant storage directory if it doesn't exist
os.makedirs(QDRANT_PATH, exist_ok=True)

# Initialize Qdrant client with the configured path
qdrant_client = None

def _initialize_qdrant_client():
    global qdrant_client
    if qdrant_client is not None:
        return qdrant_client
    
    try:
        qdrant_client = QdrantClient(path=QDRANT_PATH)
        print(f"Successfully initialized Qdrant client with path: {QDRANT_PATH}")
        
        # Disable noisy destructor that triggers portalocker ImportError at shutdown
        try:
            qdrant_client.__del__ = lambda self: None  # type: ignore
        except Exception:
            pass
            
        return qdrant_client
    except Exception as e:
        print(f"Error initializing Qdrant client: {e}")
        # If it's a lock error, try to wait and retry once
        if "already accessed by another instance" in str(e):
            import time
            print("Waiting 2 seconds and retrying Qdrant initialization...")
            time.sleep(2)
            try:
                qdrant_client = QdrantClient(path=QDRANT_PATH)
                print(f"Successfully initialized Qdrant client on retry with path: {QDRANT_PATH}")
                try:
                    qdrant_client.__del__ = lambda self: None  # type: ignore
                except Exception:
                    pass
                return qdrant_client
            except Exception as retry_e:
                print(f"Retry failed: {retry_e}")
                qdrant_client = None
        else:
            qdrant_client = None
        return None

# Try initial initialization
if not DEFER_STARTUP_INIT and not IS_RELOAD_SUPERVISOR:
    _initialize_qdrant_client()


def _safe_close_qdrant():
    """Close Qdrant client defensively before interpreter teardown."""
    global qdrant_client
    if qdrant_client is None:
        return
    try:
        qdrant_client.close()
        print("[INFO] Qdrant client closed cleanly (atexit)")
    except Exception as e:
        print(f"[DEBUG] Ignored Qdrant close issue during shutdown: {e}")

atexit.register(_safe_close_qdrant)

# Embedding model setup (local path to avoid SSL). Set EMBED_MODEL_PATH env var to a local folder.
# Use a relative path as default fallback instead of hardcoded absolute path
default_model_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "models", "all-MiniLM-L6-v2")
model_path = os.environ.get("EMBED_MODEL_PATH", default_model_path)
print(f"Attempting to load model from: {model_path}" if not IS_RELOAD_SUPERVISOR else "[SUPERVISOR] Skipping model load (deferred)")
print(f"Path exists: {os.path.exists(model_path)}" if not IS_RELOAD_SUPERVISOR else "")
print(f"Path is absolute: {os.path.isabs(model_path)}" if not IS_RELOAD_SUPERVISOR else "")

try:
    if not DEFER_STARTUP_INIT:  # Only load at import time if not deferring
        # Try to load the model from local path
        embedder = SentenceTransformer(model_path)
        print(f"Successfully loaded local embedding model from: {model_path}")
    else:
        embedder = None  # Will be loaded in lifespan
except Exception as e:
    print(f"Warning: Could not load sentence transformer model from '{model_path}': {e}")
    print("Using simple keyword-based search only")
    embedder = None

# ---- New chunking helpers ----
CHUNK_SIZE = 800
CHUNK_OVERLAP = 150
EMBEDDING_BATCH_SIZE = int(os.environ.get('EMBEDDING_BATCH_SIZE', '32'))  # NEW: used in batch embedding
# Phase 1 additions
SCHEMA_VERSION = int(os.environ.get("SCHEMA_VERSION", "3"))  # Bumped schema version for per-page + patterns + normalization
ENABLE_PYMUPDF = os.environ.get("ENABLE_PYMUPDF", "1") == "1"
AUTO_INGEST = os.environ.get("AUTO_INGEST", "1") == "1"  # If 0, skip automatic ingest on startup
FORCE_REINDEX = os.environ.get("FORCE_REINDEX", "0") == "1"  # Force full reindex
TARGET_TOKENS = 200
MAX_TOKENS = 280
HEADER_FOOTER_FREQ_THRESHOLD = 0.6
ENABLE_TABLES = os.environ.get("ENABLE_TABLES", "1") == "1"
# New: allow skipping expensive PDF re-parse on unchanged corpus
FAST_SKIP_PDF_REPARSE = os.environ.get("FAST_SKIP_PDF_REPARSE", "1") == "1"
MAX_PYPDF2_WARNINGS_PER_FILE = int(os.environ.get("MAX_PYPDF2_WARNINGS_PER_FILE", "5"))

try:
    if ENABLE_PYMUPDF:
        import fitz  # PyMuPDF
        _PYMUPDF_AVAILABLE = True
    else:
        _PYMUPDF_AVAILABLE = False
except Exception:
    _PYMUPDF_AVAILABLE = False
    fitz = None  # type: ignore
    print("[INFO] PyMuPDF not available; install with 'pip install pymupdf' to enable structured parsing")

def extract_pdf_pages(fpath: str) -> List[str]:
    """Extract text per page from a PDF using PyPDF2 with per-page pdfplumber fallback.
    Returns list of page texts. Throttles noisy warnings and retries empty pages individually.
    """
    pages_text: List[str] = []
    warn_count = 0
    try:
        with open(fpath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text() or ""
                    page_text = page_text.strip()
                    pages_text.append(page_text)
                except Exception as page_error:
                    warn_count += 1
                    if warn_count <= MAX_PYPDF2_WARNINGS_PER_FILE:
                        print(f"[WARNING] PyPDF2 failed on page {page_num + 1} of {fpath}: {page_error}")
                    elif warn_count == MAX_PYPDF2_WARNINGS_PER_FILE + 1:
                        print(f"[WARNING] Further PyPDF2 page parse warnings for {fpath} suppressed...")
                    pages_text.append("")
        # Per-page fallback for empty pages (rather than whole-file threshold only)
        if any(not p.strip() for p in pages_text):
            try:
                with pdfplumber.open(fpath) as pdf:
                    page_count = min(len(pdf.pages), len(pages_text))
                    for idx in range(page_count):
                        if pages_text[idx].strip():
                            continue  # already have text
                        try:
                            txt = pdf.pages[idx].extract_text() or ""
                            pages_text[idx] = txt.strip()
                        except Exception as fallback_page_err:
                            # Only log first few
                            if warn_count < MAX_PYPDF2_WARNINGS_PER_FILE:
                                print(f"[WARNING] pdfplumber fallback failed page {idx+1} {fpath}: {fallback_page_err}")
            except Exception as per_page_fb_err:
                # Silent unless debug desired
                print(f"[INFO] Per-page pdfplumber fallback unavailable for {fpath}: {per_page_fb_err}")
        # Legacy heuristic: if overall still minimal content, attempt full pdfplumber extraction
        if sum(len(p) for p in pages_text) < 50:
            try:
                print(f"[INFO] Low aggregate text from PyPDF2 for {fpath}, trying full pdfplumber fallback")
                pages_text_full: List[str] = []
                with pdfplumber.open(fpath) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text() or ""
                            pages_text_full.append(page_text.strip())
                        except Exception as page_error:
                            if page_num < MAX_PYPDF2_WARNINGS_PER_FILE:
                                print(f"[WARNING] pdfplumber failed on page {page_num + 1} of {fpath}: {page_error}")
                            pages_text_full.append("")
                pages_text = pages_text_full
            except Exception as whole_fb_err:
                print(f"[INFO] Full pdfplumber fallback failed for {fpath}: {whole_fb_err}")
    except Exception as pdf_error:
        print(f"[WARNING] PyPDF2 completely failed on {fpath}: {pdf_error}, trying pdfplumber")
        try:
            with pdfplumber.open(fpath) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text() or ""
                        page_text = page_text.strip()
                        pages_text.append(page_text)
                    except Exception as page_error:
                        if page_num < MAX_PYPDF2_WARNINGS_PER_FILE:
                            print(f"[WARNING] pdfplumber failed on page {page_num + 1} of {fpath}: {page_error}")
                        pages_text.append("")
        except Exception as fallback_error:
            print(f"[ERROR] Both PyPDF2 and pdfplumber failed on {fpath}: {fallback_error}")
    if not any(p.strip() for p in pages_text):
        print(f"[WARNING] No content extracted from {fpath}")
        return []
    return pages_text


def normalize_block_text(text: str) -> str:
    """Basic normalization: NFKC, collapse whitespace, trim."""
    if not text:
        return ""
    t = unicodedata.normalize("NFKC", text)
    t = ' '.join(t.split())
    return t.strip()

def chunk_text(pages: List[str]) -> List[dict]:
    """Chunk list of page texts into overlapping segments. Returns list of chunk dicts."""
    chunks: List[dict] = []
    for page_idx, page_text in enumerate(pages):
        text = page_text.strip()
        if not text:
            continue
        start = 0
        length = len(text)
        chunk_number = 0
        while start < length:
            end = min(start + CHUNK_SIZE, length)
            chunk_str = text[start:end]
            # Extend to sentence boundary minimal (look ahead for next period within 40 chars)
            if end < length:
                next_period = text.find('.', end, min(end + 40, length))
                if next_period != -1:
                    chunk_str = text[start:next_period + 1]
                    end = next_period + 1
            md5 = hashlib.md5(chunk_str.encode('utf-8')).hexdigest()
            chunks.append({
                "page": page_idx + 1,
                "pages": [page_idx + 1],
                "chunk_index": chunk_number,
                "text": chunk_str,
                "md5": md5
            })
            chunk_number += 1
            if end >= length:
                break
            start = max(end - CHUNK_OVERLAP, 0)
    return chunks

# ---- End chunking helpers ----


# Manifest for incremental ingest
INGEST_MANIFEST_PATH = os.path.join(os.path.dirname(__file__), "ingest_manifest.json")
MANIFEST_BACKUPS = 3

# Track skipped files
ingest_skipped_files = 0
# Track removed (deleted from disk) files
ingest_removed_files = 0
# Additional ingest metric
ingest_pages_reindexed = 0

# Helper: delete all vector points for a given relative filename

def delete_file_points(rel_filename: str):
    if qdrant_client is None:
        print(f"[WARNING] Cannot delete vectors for {rel_filename}: Qdrant client not available")
        return
    try:
        from qdrant_client.http import models as rest_models
        qdrant_client.delete(
            collection_name=QDRANT_COLLECTION,
            points_selector=qdrant_models.Filter(
                must=[
                    qdrant_models.FieldCondition(
                        key="filename",
                        match=qdrant_models.MatchValue(value=rel_filename)
                    )
                ]
            )
        )
        print(f"[INFO] Deleted existing vectors for {rel_filename}")
    except Exception as e:
        print(f"[WARNING] Failed to delete vectors for {rel_filename}: {e}")

def delete_file_pages_points(rel_filename: str, pages: List[int]):
    """Delete only vectors for specific pages of a file."""
    if not pages or qdrant_client is None:
        return
    try:
        for p in pages:
            qdrant_client.delete(
                collection_name=QDRANT_COLLECTION,
                points_selector=qdrant_models.Filter(
                    must=[
                        qdrant_models.FieldCondition(key="filename", match=qdrant_models.MatchValue(value=rel_filename)),
                        qdrant_models.FieldCondition(key="pages", match=qdrant_models.MatchValue(value=p))
                    ]
                )
            )
        print(f"[INFO] Deleted vectors for {rel_filename} pages {pages}")
    except Exception as e:
        print(f"[WARNING] Failed selective page delete for {rel_filename}: {e}")


def _load_manifest():
    if os.path.exists(INGEST_MANIFEST_PATH):
        try:
            with open(INGEST_MANIFEST_PATH, 'r', encoding='utf-8') as mf:
                return json.load(mf)
        except Exception as e:
            print(f"[WARNING] Could not read manifest: {e}")
    return {"files": {}}


def _save_manifest(manifest: dict):
    try:
        # Rotate backups
        if os.path.exists(INGEST_MANIFEST_PATH):
            for i in range(MANIFEST_BACKUPS-1, 0, -1):
                older = f"{INGEST_MANIFEST_PATH}.bak{i}"
                newer = f"{INGEST_MANIFEST_PATH}.bak{i-1}" if i>1 else f"{INGEST_MANIFEST_PATH}.bak1"
                if os.path.exists(newer):
                    os.replace(newer, older)
            # Copy current to .bak1
            import shutil
            shutil.copy2(INGEST_MANIFEST_PATH, f"{INGEST_MANIFEST_PATH}.bak1")
        with open(INGEST_MANIFEST_PATH, 'w', encoding='utf-8') as mf:
            json.dump(manifest, mf, indent=2)
    except Exception as e:
        print(f"[WARNING] Could not write manifest: {e}")


def load_documents():
    global documents, ingest_total_files, ingest_processed_files, ingest_skipped_files, ingest_removed_files, ingest_pages_reindexed
    documents = []
    ingest_processed_files = 0
    ingest_skipped_files = 0
    ingest_removed_files = 0
    ingest_pages_reindexed = 0
    manifest = _load_manifest()
    files_section = manifest.get('files', {})
    file_patterns = ["**/*.txt", "**/*.pdf", "**/*.docx", "**/*.xls", "**/*.xlsx"]
    all_files = []
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
        ingest_total_files = 0
        return
    for pattern in file_patterns:
        all_files.extend(glob.glob(os.path.join(DATA_DIR, pattern), recursive=True))
    ingest_total_files = len(all_files)
    if embedder is not None and qdrant_client is not None and QDRANT_COLLECTION not in [c.name for c in qdrant_client.get_collections().collections]:
        qdrant_client.create_collection(
            collection_name=QDRANT_COLLECTION,
            vectors_config=qdrant_models.VectorParams(size=384, distance=qdrant_models.Distance.COSINE)
        )
    elif qdrant_client is None:
        print("[WARNING] Qdrant client not available - vector operations will be skipped")
    elif embedder is None:
        print("[WARNING] Embedder not available - vector operations will be skipped")
    updated_manifest_files = files_section

    # Early exit if force reindex not requested and all hashes + schema match (schema optional for legacy entries)
    if not FORCE_REINDEX and updated_manifest_files and all_files:
        all_unchanged = True
        for fpath in all_files:
            rel = os.path.relpath(fpath, DATA_DIR)
            prev_entry = updated_manifest_files.get(rel)
            if not prev_entry:
                all_unchanged = False
                break
            prev_schema = prev_entry.get('schema_version')
            # If schema_version present and different -> must reindex. If missing, treat as legacy acceptable.
            if prev_schema is not None and prev_schema != SCHEMA_VERSION:
                all_unchanged = False
                break
            # compute file_md5
            try:
                file_md5_hash = hashlib.md5()
                with open(fpath, 'rb') as fb:
                    for chunk in iter(lambda: fb.read(1024 * 1024), b''):
                        file_md5_hash.update(chunk)
                file_md5 = file_md5_hash.hexdigest()
                if prev_entry.get('file_md5') != file_md5:
                    all_unchanged = False
                    break
            except Exception:
                all_unchanged = False
                break
        if all_unchanged:
            print("[INFO] All files unchanged – skipping embedding phase (legacy schema tolerated)")
            for fpath in all_files:
                rel = os.path.relpath(fpath, DATA_DIR)
                ext = os.path.splitext(fpath)[1].lower()
                try:
                    if ext == '.txt':
                        with open(fpath, 'r', encoding='utf-8') as f:
                            content = f.read()
                    elif ext == '.pdf':
                        if FAST_SKIP_PDF_REPARSE:
                            # Avoid triggering PyPDF2 warnings again for unchanged PDFs
                            content = f"(unchanged pdf: {rel} – content not re-parsed; set FAST_SKIP_PDF_REPARSE=0 to load text)"
                        else:
                            pages = extract_pdf_pages(fpath)
                            content = "\n".join(pages)
                    elif ext == '.docx':
                        doc = Document(fpath)
                        content = "\n".join(p.text for p in doc.paragraphs)
                    else:
                        continue
                    documents.append({"content": content, "filename": rel, "structured": "{}"})
                except Exception as e:
                    print(f"[WARNING] Failed to load content for {rel} during fast path: {e}")
            ingest_processed_files = ingest_total_files
            ingest_skipped_files = ingest_total_files
            return
    current_rel_set = set()
    # Normal processing path
    for fpath in all_files:
        rel = os.path.relpath(fpath, DATA_DIR)
        current_rel_set.add(rel)
        ext = os.path.splitext(fpath)[1].lower()
        try:
            file_md5_hash = hashlib.md5()
            with open(fpath, 'rb') as fb:
                for chunk in iter(lambda: fb.read(1024 * 1024), b''):
                    file_md5_hash.update(chunk)
            file_md5 = file_md5_hash.hexdigest()
            prev_entry = updated_manifest_files.get(rel)
            prev_pages_meta = prev_entry.get('pages') if prev_entry else []
            prev_page_map = {p['page']: p for p in prev_pages_meta} if prev_pages_meta else {}
            prev_schema = prev_entry.get('schema_version') if prev_entry else None
            unchanged_whole = (
                not FORCE_REINDEX and
                prev_entry and
                prev_entry.get('file_md5') == file_md5 and
                (prev_schema is None or prev_schema == SCHEMA_VERSION)
            )
            parse_needed = True
            pages: List[str] = []
            structured_pages_blocks = None
            page_md5s: List[str] = []
            per_page_num_blocks: List[int] = []
            header_patterns: List[str] = []
            footer_patterns: List[str] = []
            block_md5s: List[str] = []
            # Extraction
            if ext == '.txt':
                with open(fpath, 'r', encoding='utf-8') as f:
                    txt = f.read()
                pages = [txt]
                norm = normalize_block_text(txt)
                page_md5s = [hashlib.md5(norm.encode('utf-8')).hexdigest()]
                per_page_num_blocks = [1 if norm else 0]
                block_md5s = [hashlib.md5(norm.encode('utf-8')).hexdigest()]
            elif ext == '.pdf':
                structured_result = _structured_parse_pdf(fpath) if _PYMUPDF_AVAILABLE else None
                if structured_result:
                    structured_pages_blocks = structured_result['pages_blocks']
                    pages = structured_result['page_texts']
                    header_patterns = structured_result['header_lines']
                    footer_patterns = structured_result['footer_lines']
                    per_page_num_blocks = structured_result['per_page_num_blocks']
                    block_md5s = structured_result['block_md5s']
                    page_md5s = [hashlib.md5(normalize_block_text(p).encode('utf-8')).hexdigest() for p in pages]
                else:
                    pages = extract_pdf_pages(fpath)
                    page_md5s = [hashlib.md5(normalize_block_text(p).encode('utf-8')).hexdigest() for p in pages]
                    per_page_num_blocks = [1 if p.strip() else 0 for p in pages]
                    block_md5s = [hashlib.md5(normalize_block_text(p).encode('utf-8')).hexdigest() for p in pages]
                if not pages:
                    ingest_processed_files += 1
                    continue
            elif ext == '.docx':
                doc = Document(fpath)
                doc_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                pages = [doc_text]
                norm = normalize_block_text(doc_text)
                page_md5s = [hashlib.md5(norm.encode('utf-8')).hexdigest()]
                per_page_num_blocks = [1 if norm else 0]
                block_md5s = [hashlib.md5(norm.encode('utf-8')).hexdigest()]
            else:
                ingest_processed_files += 1
                continue
            # Determine selective page changes
            changed_pages: List[int] = []
            if prev_page_map and prev_entry and prev_entry.get('schema_version') == SCHEMA_VERSION:
                for idx, pg_md5 in enumerate(page_md5s, start=1):
                    prev_pg = prev_page_map.get(idx)
                    if not prev_pg or prev_pg.get('page_md5') != pg_md5:
                        changed_pages.append(idx)
            else:
                # No prior per-page info -> treat all as changed if file changed
                changed_pages = list(range(1, len(page_md5s)+1)) if not unchanged_whole else []
            # If whole file unchanged, skip
            if unchanged_whole:
                ingest_skipped_files += 1
                ingest_processed_files += 1
                updated_manifest_files[rel] = prev_entry  # retain existing rich manifest
                print(f"[INFO] Skipped unchanged file (schema+hash match): {rel}")
                content = "\n".join(pages)
                structured_doc_meta = {"filename": os.path.basename(fpath), "content_length": len(content), "file_type": ext}
                documents.append({"content": content, "filename": rel, "structured": str(structured_doc_meta)})
                continue
            # Partial reindex if some pages changed but not all
            partial_reindex = prev_entry is not None and changed_pages and len(changed_pages) < len(page_md5s)
            if partial_reindex:
                delete_file_pages_points(rel, changed_pages)
                ingest_pages_reindexed += len(changed_pages)
            elif prev_entry:  # full reindex
                delete_file_points(rel)
            content = "\n".join(pages)
            first_line = content.split('\n')[0] if content else ""
            structured_doc_meta = {"filename": os.path.basename(fpath), "first_line": first_line[:100], "content_length": len(content), "file_type": ext}
            documents.append({"content": content, "filename": rel, "structured": str(structured_doc_meta)})
            # Chunking
            if structured_pages_blocks:
                chunk_list = _structure_aware_chunk(structured_pages_blocks)
            else:
                chunk_list = chunk_text(pages)
            # Filter chunks for partial reindex
            if partial_reindex:
                changed_set = set(changed_pages)
                chunk_list_to_embed = [c for c in chunk_list if any(p in changed_set for p in (c.get('pages') or [c.get('page')]))]
            else:
                chunk_list_to_embed = chunk_list
            chunk_md5s_all_previous = set(prev_entry.get('chunk_md5s', [])) if prev_entry else set()
            new_chunk_md5s: List[str] = []
            if embedder is not None and qdrant_client is not None and chunk_list_to_embed:
                points = []
                # NEW: batch embedding for efficiency
                try:
                    texts_to_embed = [ch["text"] for ch in chunk_list_to_embed]
                    embeddings = embedder.encode(texts_to_embed, batch_size=EMBEDDING_BATCH_SIZE).tolist() if embedder is not None else []
                except Exception as be:
                    print(f"[WARNING] Batch embedding failed, falling back to per-chunk: {be}")
                    embeddings = None
                for idx_ch, ch in enumerate(chunk_list_to_embed):
                    try:
                        if embeddings is not None:
                            embedding = embeddings[idx_ch]
                        else:
                            embedding = embedder.encode(ch["text"]).tolist()
                        base_pages = ch.get('pages') or [ch.get('page')]
                        pages_key = ','.join(str(p) for p in base_pages)
                        chunk_uuid_seed = f"{fpath}#pages={pages_key}#chunk={ch['chunk_index']}#md5={ch['md5']}#schema={SCHEMA_VERSION}"
                        point_id = str(uuid.uuid5(uuid.NAMESPACE_URL, chunk_uuid_seed))
                        payload = {
                            "filename": rel,
                            "page": ch.get("page"),
                            "pages": base_pages,
                            "page_range": ch.get("page_range"),
                            "chunk_index": ch["chunk_index"],
                            "text": ch["text"],
                            "md5": ch["md5"],
                            "structured": str(structured_doc_meta),
                            "schema_version": SCHEMA_VERSION,
                        }
                        if 'heading_path' in ch:
                            payload['heading_path'] = ch['heading_path']
                        if 'block_types' in ch:
                            payload['block_types'] = ch['block_types']
                        if 'block_types' in ch and len(ch['block_types']) == 1 and ch['block_types'][0] == 'table':
                            payload['table_markdown'] = ch['text'][:5000]
                            if '\n' in ch['text']:
                                first_line = ch['text'].split('\n',1)[0]
                                if 'rows=' in first_line and 'cols=' in first_line:
                                    payload['table_summary'] = first_line[:300]
                            if 'table_md5' in ch:
                                payload['table_md5'] = ch['table_md5']
                            if 'table_csv' in ch:
                                payload['table_csv'] = ch['table_csv'][:TABLE_CSV_MAX_CHARS]
                            if 'n_rows' in ch:
                                payload['n_rows'] = ch['n_rows']
                            if 'n_cols' in ch:
                                payload['n_cols'] = ch['n_cols']
                        points.append(qdrant_models.PointStruct(id=point_id, vector=embedding, payload=payload))
                        new_chunk_md5s.append(ch['md5'])
                    except Exception as embed_error:
                        print(f"[WARNING] Could not embed chunk {ch.get('chunk_index')} of {rel}: {embed_error}")
                if points:
                    try:
                        qdrant_client.upsert(collection_name=QDRANT_COLLECTION, points=points)
                        action = "partially reindexed" if partial_reindex else "indexed"
                        print(f"[INFO] {action} {len(points)} chunks for {rel} (schema {SCHEMA_VERSION}) [batch]")
                    except Exception as up_err:
                        print(f"[ERROR] Upsert failed for {fpath}: {up_err}")
            elif embedder is None:
                print(f"[INFO] Skipping vector storage for {rel} (no embedder available)")
            elif qdrant_client is None:
                print(f"[INFO] Skipping vector storage for {rel} (Qdrant client not available)")
            else:
                print(f"[WARNING] No chunks produced for {rel}")
            # Merge chunk md5s
            final_chunk_md5s = list(chunk_md5s_all_previous.union(set(new_chunk_md5s))) if partial_reindex else new_chunk_md5s
            pages_meta = [{"page": i+1, "page_md5": page_md5s[i], "num_blocks": per_page_num_blocks[i], "ocr": False} for i in range(len(page_md5s))]
            block_index_hash = hashlib.md5(''.join(block_md5s).encode('utf-8')).hexdigest() if block_md5s else ''
            tables_count = 0
            if structured_pages_blocks and 'table' in [b.get('type') for page_b in structured_pages_blocks for b in page_b]:
                tables_count = sum(1 for page_b in structured_pages_blocks for b in page_b if b.get('type') == 'table')
            updated_manifest_files[rel] = {
                "file_md5": file_md5,
                "num_chunks": len(final_chunk_md5s),
                "chunk_md5s": final_chunk_md5s,
                "schema_version": SCHEMA_VERSION,
                "pages": pages_meta,
                "header_patterns": header_patterns,
                "footer_patterns": footer_patterns,
                "block_index_hash": block_index_hash,
                "counts": {"tables": tables_count, "equations": 0, "figures": 0, "ocr_pages": 0, "chars": sum(len(p) for p in pages)}
            }
        except Exception as e:
            print(f"[ERROR] Failed to process {fpath}: {e}")
        ingest_processed_files += 1

    # Detect removed files
    previous_rel_set = set(updated_manifest_files.keys())
    removed = previous_rel_set - current_rel_set
    if removed:
        for rel in removed:
            delete_file_points(rel)
            updated_manifest_files.pop(rel, None)
            ingest_removed_files += 1
            print(f"[INFO] Removed stale file entry: {rel}")

    manifest['files'] = updated_manifest_files
    _save_manifest(manifest)
    _build_sparse_index()


def simple_search(query: str, docs: List[dict], top_k: int = 3) -> List[dict]:
    """Simple keyword-based search"""
    query_words = query.lower().split()
    scored_docs = []
    
    for doc in docs:
        content_lower = doc["content"].lower()
        score = sum(1 for word in query_words if word in content_lower)
        if score > 0:
            scored_docs.append({"doc": doc, "score": score})
    
    # Sort by score and return top_k
    scored_docs.sort(key=lambda x: x["score"], reverse=True)
    return [item["doc"] for item in scored_docs[:top_k]]


def vector_search(query: str, top_k: int = 3) -> List[dict]:
    if embedder is None or qdrant_client is None:
        return []
    try:
        query_vec = embedder.encode(query).tolist()
        results = []
        try:
            # New preferred API
            qp = qdrant_client.query_points(
                collection_name=QDRANT_COLLECTION,
                query=query_vec,
                limit=top_k,
                with_payload=True
            )
            hits = qp.points
        except Exception as new_api_err:
            print(f"[WARNING] query_points failed, falling back to search: {new_api_err}")
            hits = qdrant_client.search(
                collection_name=QDRANT_COLLECTION,
                query_vector=query_vec,
                limit=top_k,
                with_payload=True
            )
        for hit in hits:
            payload = hit.payload or {}
            results.append({
                "filename": payload.get("filename", "unknown"),
                "content": payload.get("text", payload.get("content", "")),
                "page": payload.get("page"),
                "chunk_index": payload.get("chunk_index")
            })
        return results
    except Exception as e:
        print(f"[ERROR] Vector search failed: {e}")
        return []


def generate_answer(question: str, context_docs: List[dict]) -> str:
    """Answer generation using Llama3 via Ollama API with adaptive timeout and retry"""
    print(f"[DEBUG] Received question: {question}")
    print(f"[DEBUG] Context docs: {[doc['filename'] for doc in context_docs]}")
    if not context_docs:
        print("[DEBUG] No relevant documents found.")
        return "I couldn't find relevant information to answer your question."

    # Optimize context by limiting the number of documents and snippet length
    max_docs = min(len(context_docs), 8)  # Limit to 8 most relevant documents
    snippet_length = 300  # Initial snippet length
    
    # Adaptive timeout and retry mechanism
    max_retries = 3  # Increased from 2 to 3
    initial_timeout = 90  # Initial timeout in seconds
    min_timeout = 30  # Minimum timeout for subsequent retries
    timeout_reduction_factor = 0.7  # Reduce timeout by 30% on each retry
    
    # Track timeouts to reduce context size on multiple timeouts
    timeout_count = 0
    
    for attempt in range(max_retries + 1):
        # Adjust context size if we've had multiple timeouts
        current_max_docs = max(4, max_docs - timeout_count)  # Reduce docs but keep at least 4
        current_snippet_length = max(150, snippet_length - (timeout_count * 50))  # Reduce snippet length but keep at least 150 chars
        
        # Build context with current parameters
        context_text = "\n\n".join([f"From {doc.get('filename','unknown')} (page {doc.get('page','?')}):\n{doc.get('content','')[:current_snippet_length]}..." 
                                for doc in context_docs[:current_max_docs]])
        
        # More conversational system prompt
        system_prompt = (
            "You are a knowledgeable tutor and academic guide. "
            "Use the provided context and, when helpful, your broader knowledge to teach the student clearly. "
            "Provide concise explanations, step-by-step reasoning, and encourage learning with follow-up questions."
        )
        prompt = f"{system_prompt}\n\nQuestion: {question}\n\nContext:\n{context_text}"
        
        # Calculate current timeout (reduce on each retry)
        current_timeout = max(min_timeout, initial_timeout * (timeout_reduction_factor ** attempt))
        
        print(f"[INFO] Attempt {attempt+1}/{max_retries+1}: Using {current_max_docs} docs, {current_snippet_length} chars per snippet, {current_timeout:.1f}s timeout")
        print(f"[DEBUG] Sending prompt to Ollama: {prompt[:200]}...")
        
        try:
            print(f"[INFO] Sending request to Ollama (attempt {attempt+1}/{max_retries+1})")
            response = requests.post(
                "http://localhost:11434/api/generate",
                json={
                    "model": "llama3:latest",
                    "prompt": prompt,
                    "stream": False,
                    "temperature": 0.7  # Add temperature parameter to potentially speed up generation
                },
                timeout=current_timeout
            )
            print(f"[DEBUG] Ollama response status: {response.status_code}")
            response.raise_for_status()
            data = response.json()
            answer = data.get("response", "")
            
            # Check for empty response
            if not answer.strip():
                print(f"[WARNING] Empty response received on attempt {attempt+1}")
                if attempt < max_retries:
                    timeout_count += 1  # Treat empty response like a timeout
                    continue
                else:
                    return "Error: Received empty response from Ollama. Please try again with a simpler question."
            
            print(f"[DEBUG] Final answer: {answer[:200]}...")
            print(f"[INFO] Ollama response received successfully on attempt {attempt+1}")
            return answer
            
        except requests.exceptions.ConnectionError as e:
            print(f"[ERROR] Connection error on attempt {attempt+1}: {e}")
            # For connection errors, we might need to wait longer for Ollama to recover
            if attempt < max_retries:
                wait_time = (attempt + 1) * 3  # Longer wait for connection issues
                print(f"[INFO] Connection issue - waiting {wait_time} seconds before retry...")
                time.sleep(wait_time)
            else:
                return "Error: Cannot connect to Ollama. Please check if Ollama is running and restart if necessary."
                
        except requests.exceptions.Timeout as e:
            print(f"[WARNING] Timeout on attempt {attempt+1}/{max_retries+1}: {e}")
            timeout_count += 1
            
            if attempt < max_retries:
                # Wait before retrying, with increasing backoff
                wait_time = (attempt + 1) * 2
                print(f"[INFO] Timeout occurred - waiting {wait_time} seconds before retry with reduced context...")
                time.sleep(wait_time)
            else:
                return f"Error: Request to Ollama timed out after {max_retries+1} attempts. Try asking a simpler question or check if Ollama is running properly."
                
        except Exception as e:
            print(f"[ERROR] Unexpected error on attempt {attempt+1}: {e}")
            if attempt < max_retries:
                wait_time = (attempt + 1) * 2
                print(f"[INFO] Error occurred - waiting {wait_time} seconds before retry...")
                time.sleep(wait_time)
            else:
                return f"Error: Unexpected issue when generating answer: {str(e)}. Please try again later."
            import traceback
            traceback.print_exc()
            return f"Error communicating with Ollama: {str(e)}"


@app.get("/ingest")
def ingest():
    load_documents()
    return {"status": "reloaded", "documents_loaded": len(documents)}

@app.get("/health")
def health_check():
    """Check if Ollama and Llama3 are available with detailed diagnostics"""
    try:
        # Use a shorter timeout for health checks
        response = requests.get("http://localhost:11434/api/tags", timeout=3)
        response.raise_for_status()
        models = response.json().get("models", [])
        llama3_available = any("llama3" in model.get("name", "") for model in models)
        
        # Check if the model is ready by sending a minimal request
        if llama3_available:
            try:
                # Send a minimal prompt to check if model is responsive
                test_response = requests.post(
                    "http://localhost:11434/api/generate",
                    json={
                        "model": "llama3:latest",
                        "prompt": "Hello",
                        "stream": False,
                        "temperature": 0.7
                    },
                    timeout=3
                )
                test_response.raise_for_status()
                
                # Check for empty response
                response_data = test_response.json()
                if not response_data.get("response"):
                    print("[WARNING] Model returned empty response")
                    model_ready = False
                    model_error = "Model returned empty response"
                else:
                    model_ready = True
                    model_error = None
            except requests.exceptions.Timeout as model_timeout:
                print(f"[WARNING] Model health check timed out: {model_timeout}")
                model_ready = False
                model_error = "Model response timed out"
            except requests.exceptions.ConnectionError as model_conn_err:
                print(f"[WARNING] Model health check connection error: {model_conn_err}")
                model_ready = False
                model_error = "Connection error during model check"
            except Exception as model_err:
                print(f"[WARNING] Model health check failed: {model_err}")
                model_ready = False
                model_error = str(model_err)
        else:
            model_ready = False
            model_error = "Llama3 model not found in available models"
            
        return {
            "ollama_running": True,
            "llama3_available": llama3_available,
            "model_ready": model_ready,
            "available_models": [model.get("name", "") for model in models],
            "model_error": model_error if not model_ready else None
        }
    except requests.exceptions.Timeout:
        return {
            "ollama_running": False,
            "llama3_available": False,
            "model_ready": False,
            "error": "Connection to Ollama timed out",
            "troubleshooting": "Check if Ollama is running but overloaded or unresponsive"
        }
    except requests.exceptions.ConnectionError:
        return {
            "ollama_running": False,
            "llama3_available": False,
            "model_ready": False,
            "error": "Ollama is not running",
            "troubleshooting": "Start Ollama service or check if it's running on a different port"
        }
    except Exception as e:
        return {
            "ollama_running": False,
            "llama3_available": False,
            "model_ready": False,
            "error": str(e),
            "troubleshooting": "Unexpected error, check Ollama installation and configuration"
        }

# Simple LRU cache for recent questions to avoid redundant processing
from functools import lru_cache

@lru_cache(maxsize=20)
def get_context_for_question(question: str, use_hybrid: bool, top_k_dense=6, top_k_sparse=6, final_k=6):
    """Cache context retrieval for similar questions with optimized parameters"""
    if use_hybrid:
        raw_docs = hybrid_search(question, top_k_dense=top_k_dense, top_k_sparse=top_k_sparse, final_k=final_k)
        return assemble_context(question, raw_docs, max_tokens=MAX_CONTEXT_TOKENS)
    elif embedder is not None and qdrant_client is not None and QDRANT_COLLECTION in [c.name for c in qdrant_client.get_collections().collections]:
        raw_docs = vector_search(question, top_k=final_k)
        return assemble_context(question, raw_docs, max_tokens=MAX_CONTEXT_TOKENS)
    else:
        print("[INFO] Using simple search fallback")
        return simple_search(question, documents, top_k=final_k)

@app.post("/ask")
async def ask(request: Request):
    print("[DEBUG] /ask endpoint called")
    start_time = time.time()
    data = await request.json()
    question = data.get("question", "")
    session_id = data.get("session_id")  # session id for history
    print(f"[DEBUG] Question received: '{question}'")
    if session_id:
        db = SessionLocal()
        try:
            cs = db.query(ChatSummary).filter(ChatSummary.session_id == session_id).first()
            chat_summary_text = cs.summary_text if cs and cs.summary_text else ""
        finally:
            db.close()
        print(f"[DEBUG] Chat summary for session {session_id}: {chat_summary_text}")
    else:
        chat_summary_text = ""

    if not question:
        print("[DEBUG] No question provided")
        return {"error": "Question is required"}
    if qdrant_client is None:
        _initialize_qdrant_client()
    use_hybrid = ENABLE_SPARSE and _SPARSE_OK and embedder is not None and qdrant_client is not None
    try:
        context_time_start = time.time()
        try:
            context_docs = get_context_for_question(question, use_hybrid)
        except Exception as context_error:
            print(f"[ERROR] Context retrieval failed: {context_error}")
            context_docs = simple_search(question, documents, top_k=5)
        context_time = time.time() - context_time_start
        print(f"[TIMING] Context retrieval took {context_time:.2f}s")
        if chat_summary_text:
            context_docs.insert(0, {'filename': 'chat_summary', 'content': chat_summary_text, 'page': 0, 'chunk_index': 0})
        # Unified generation
        answer_time_start = time.time()
        answer, citations, meta = generate_answer_engine(question, context_docs, chat_summary_text, session_id)
        answer_time = time.time() - answer_time_start
        print(f"[TIMING] Answer generation took {answer_time:.2f}s (topic_shift={meta.get('topic_shift')}, history_used={meta.get('history_used')})")
        if answer.lower().startswith("error"):
            print("[WARNING] unified engine error fallback -> extractive summary")
            summary_lines = []
            for d in context_docs[:4]:
                snippet = (d.get('content','') or '')[:120].strip().replace("\n", " ")
                summary_lines.append(f"- {d.get('filename')} (p{d.get('page')}): {snippet}...")
            answer = answer + "\n\nRelevant excerpts:\n" + "\n".join(summary_lines)
        context_for_display = []
        for doc in context_docs[:6]:
            snippet = (doc.get('content','') or '')[:120].strip()
            context_for_display.append({
                "id": doc.get('citation_id', 0),
                "filename": doc.get('filename', 'Unknown'),
                "page": doc.get('page', 0),
                "content": snippet
            })
        total_time = time.time() - start_time
        print(f"[TIMING] Total processing took {total_time:.2f}s")
        return {
            "answer": answer,
            "context": context_for_display,
            "documents_found": len(context_docs),
            "citations": citations,
            "processing_time": round(total_time, 2),
            "meta": meta
        }
    except Exception as e:
        print(f"[ERROR] /ask endpoint failure: {e}")
        return {"error": str(e)}

@app.get('/tables')
def list_tables(filename: Optional[str] = None, limit: int = 200):
    """List table chunks (single-table chunks with table_markdown). Optional filter by filename."""
    if embedder is None or qdrant_client is None:
        return {"error": "Vector store unavailable"}
    try:
        must = []
        if filename:
            must.append(qdrant_models.FieldCondition(key="filename", match=qdrant_models.MatchValue(value=filename)))
        filt = qdrant_models.Filter(must=must) if must else None
        tables = []
        offset = None
        batch_limit = 256
        while len(tables) < limit:
            points, offset = qdrant_client.scroll(collection_name=QDRANT_COLLECTION, scroll_filter=filt, limit=batch_limit, with_payload=True, offset=offset)
            if not points:
                break
            for pt in points:
                payload = pt.payload or {}
                if 'table_markdown' in payload:
                    tables.append({
                        'id': str(pt.id),
                        'filename': payload.get('filename'),
                        'page': payload.get('page'),
                        'chunk_index': payload.get('chunk_index'),
                        'table_md5': payload.get('table_md5'),
                        'n_rows': payload.get('n_rows'),
                        'n_cols': payload.get('n_cols'),
                        'table_summary': payload.get('table_summary'),
                        'preview': (payload.get('table_markdown') or '')[:250]
                    })
                    if len(tables) >= limit:
                        break
            if offset is None or len(tables) >= limit:
                break
        return {'count': len(tables), 'tables': tables, 'filename': filename}
    except Exception as e:
        return {'error': str(e)}

@app.get('/table_csv')
def get_table_csv(table_md5: Optional[str] = None, filename: Optional[str] = None, page: Optional[int] = None, chunk_index: Optional[int] = None):
    """Return CSV + markdown for a specific table. Identify via table_md5 OR (filename,page,chunk_index)."""
    if embedder is None or qdrant_client is None:
        return {"error": "Vector store unavailable"}
    if not table_md5 and not (filename and page is not None and chunk_index is not None):
        return {"error": "Provide table_md5 or (filename,page,chunk_index)"}
    try:
        # Strategy: scroll (filtered by filename if provided) and match conditions client-side.
        must = []
        if filename:
            must.append(qdrant_models.FieldCondition(key="filename", match=qdrant_models.MatchValue(value=filename)))
        filt = qdrant_models.Filter(must=must) if must else None
        offset = None
        batch_limit = 256
        while True:
            points, offset = qdrant_client.scroll(collection_name=QDRANT_COLLECTION, scroll_filter=filt, limit=batch_limit, with_payload=True, offset=offset)
            if not points:
                break
            for pt in points:
                payload = pt.payload or {}
                if 'table_markdown' not in payload:
                    continue
                if table_md5 and payload.get('table_md5') != table_md5:
                    continue
                if not table_md5 and filename and (payload.get('page') != page or payload.get('chunk_index') != chunk_index):
                    continue
                if not table_md5 and not filename:
                    continue
                return {
                    'id': str(pt.id),
                    'filename': payload.get('filename'),
                    'page': payload.get('page'),
                    'chunk_index': payload.get('chunk_index'),
                    'table_md5': payload.get('table_md5'),
                    'n_rows': payload.get('n_rows'),
                    'n_cols': payload.get('n_cols'),
                    'table_summary': payload.get('table_summary'),
                    'table_markdown': payload.get('table_markdown'),
                    'table_csv': payload.get('table_csv')
                }
            if offset is None:
                break
        return {'error': 'Table not found'}
    except Exception as e:
        return {'error': str(e)}


# Chat summary DB setup
CHAT_DB_PATH = os.path.join(os.path.dirname(__file__), "chat_summary.db")
engine = create_engine(f"sqlite:///{CHAT_DB_PATH}", connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
Base = declarative_base()

class ChatSummary(Base):
    __tablename__ = "chat_summaries"
    id = Column(Integer, primary_key=True, index=True)
    session_id = Column(String, index=True)
    timestamp = Column(DateTime, default=datetime.utcnow)
    mood = Column(String)
    motivation_level = Column(String)
    understanding_level = Column(String)
    summary_text = Column(String)

Base.metadata.create_all(bind=engine)

# Pydantic models for chat summary endpoint
class ChatSummaryRequest(BaseModel):
    session_id: str
    messages: List[str]

class ChatSummaryResponse(BaseModel):
    mood: str
    motivation_level: str
    understanding_level: str
    summary: str
    

@app.post("/chat_summary", response_model=ChatSummaryResponse)
def create_chat_summary(request: ChatSummaryRequest):
    """Create a new chat summary entry."""
    db = SessionLocal()
    try:
        summary = ChatSummary(
            session_id=request.session_id,
            mood="",
            motivation_level="",
            understanding_level="",
            summary_text=""
        )
        db.add(summary)
        db.commit()
        db.refresh(summary)
        return summary
    except Exception as e:
        print(f"[ERROR] Failed to create chat summary: {e}")
        return {"error": str(e)}
    finally:
        db.close()

@app.get("/chat_summary")
def read_chat_summaries(skip: int = 0, limit: int = 10):
    """Retrieve chat summaries with pagination."""
    db = SessionLocal()
    try:
        summaries = db.query(ChatSummary).offset(skip).limit(limit).all()
        return summaries
    except Exception as e:
        print(f"[ERROR] Failed to read chat summaries: {e}")
        return {"error": str(e)}
    finally:
        db.close()
    

@app.get("/chat_summary/{session_id}")
def get_chat_summary(session_id: str):
    """Get a specific chat summary by session ID."""
    db = SessionLocal()
    try:
        summary = db.query(ChatSummary).filter(ChatSummary.session_id == session_id).first()
        if summary:
            return summary
        else:
            return {"error": "Summary not found"}
    except Exception as e:
        print(f"[ERROR] Failed to get chat summary: {e}")
        return {"error": str(e)}
    finally:
        db.close()
    

@app.put("/chat_summary/{session_id}")
def update_chat_summary(session_id: str, request: ChatSummaryRequest):
    """Update an existing chat summary."""
    db = SessionLocal()
    try:
        summary = db.query(ChatSummary).filter(ChatSummary.session_id == session_id)
        if summary.first() is None:
            return {"error": "Summary not found"}
        summary.update({
            "mood": request.mood,
            "motivation_level": request.motivation_level,
            "understanding_level": request.understanding_level,
            "summary_text": request.summary
        })
        db.commit()
        return {"status": "updated"}
    except Exception as e:
        print(f"[ERROR] Failed to update chat summary: {e}")
        return {"error": str(e)}
    finally:
        db.close()
    

@app.delete("/chat_summary/{session_id}")
def delete_chat_summary(session_id: str):
    """Delete a chat summary."""
    db = SessionLocal()
    try:
        result = db.query(ChatSummary).filter(ChatSummary.session_id == session_id).delete()
        db.commit()
        if result:
            return {"status": "deleted"}
        else:
            return {"error": "Summary not found"}
    except Exception as e:
        print(f"[ERROR] Failed to delete chat summary: {e}")
        return {"error": str(e)}
    finally:
        db.close()


if __name__ == '__main__':
    import uvicorn
    import socket
    host = os.environ.get('HOST', '0.0.0.0')
    port = int(os.environ.get('PORT', '8000'))
    reload_flag = os.environ.get('RELOAD','1')=='1'
    # Auto-disable reload if using embedded Qdrant path to avoid lock contention
    if reload_flag and FORCE_DISABLE_RELOAD_FOR_EMBEDDED and not os.environ.get('QDRANT_URL'):
        print('[INFO] Auto-disabling uvicorn reload because embedded Qdrant does not support multi-process access. Set FORCE_DISABLE_RELOAD_FOR_EMBEDDED=0 to override.')
        reload_flag = False
    # Port availability check (Windows)
    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    try:
        sock.bind((host, port))
        sock.close()
    except OSError as e:
        print(f'[ERROR] Port {port} is already in use. Please stop any other server using this port and try again.')
        import sys
        sys.exit(1)
    uvicorn.run('main:app', host=host, port=port, reload=reload_flag)